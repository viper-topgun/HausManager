"""
Witter Heizkostenabrechnung – DOCX and PDF Export Service

Generates a structured billing document (Abrechnung) for the Witter form
in both Word (.docx) and PDF format.
"""

from io import BytesIO


# ─── Formatting Helpers ───────────────────────────────────────────────────────

def _n(val, dec: int = 2) -> str:
    """Format number German-style: 1.234,56"""
    if val is None:
        return ""
    try:
        s = f"{float(val):,.{dec}f}"
        return s.replace(",", "\x00").replace(".", ",").replace("\x00", ".")
    except (TypeError, ValueError):
        return ""


def _d(val) -> str:
    """Convert YYYY-MM-DD → DD.MM.YYYY"""
    if not val:
        return ""
    p = str(val).split("-")
    return f"{p[2]}.{p[1]}.{p[0]}" if len(p) == 3 else str(val)


def _s(val) -> str:
    return str(val) if val is not None else ""


# ─── DOCX Export ──────────────────────────────────────────────────────────────

def generate_docx(data: dict, year: int) -> BytesIO:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    PRIMARY_HEX = "1A7A5E"
    PRIMARY_RGB = RGBColor(0x1A, 0x7A, 0x5E)
    WHITE_RGB   = RGBColor(0xFF, 0xFF, 0xFF)
    ALT_HEX     = "E6F4EF"

    doc = Document()

    # A4 page, 2 cm margins
    sec = doc.sections[0]
    sec.page_width  = Cm(21)
    sec.page_height = Cm(29.7)
    for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, attr, Cm(2))

    # ── Helpers ──────────────────────────────────────────────────────────────

    def set_cell_bg(cell, hex_color: str):
        tcPr = cell._tc.get_or_add_tcPr()
        for existing in tcPr.findall(qn("w:shd")):
            tcPr.remove(existing)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color)
        tcPr.append(shd)

    def add_section_heading(text: str):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = PRIMARY_RGB
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    "4")
        bot.set(qn("w:color"), PRIMARY_HEX)
        pBdr.append(bot)
        pPr.append(pBdr)

    def add_table(headers: list, rows: list, widths: list):
        tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
        tbl.style    = "Table Grid"
        tbl.autofit  = False

        # Header row
        for i, (cell, h) in enumerate(zip(tbl.rows[0].cells, headers)):
            cell.text = h.replace("\n", " ")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if p.runs:
                r = p.runs[0]
                r.bold = True
                r.font.size = Pt(8)
                r.font.color.rgb = WHITE_RGB
            set_cell_bg(cell, PRIMARY_HEX)

        # Data rows
        for r_idx, row_vals in enumerate(rows):
            for c_idx, (cell, val) in enumerate(zip(tbl.rows[r_idx + 1].cells, row_vals)):
                cell.text = val
                if cell.paragraphs and cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].font.size = Pt(8)
                if r_idx % 2 == 1:
                    set_cell_bg(cell, ALT_HEX)

        # Column widths
        for row in tbl.rows:
            for i, cell in enumerate(row.cells):
                if i < len(widths):
                    cell.width = widths[i]

    def spacer():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)

    # ── Title ─────────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Witter Heizkostenabrechnung {year}")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = PRIMARY_RGB

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kd  = data.get("kundennummer", "")
    von = _d(data.get("von", ""))
    bis = _d(data.get("bis", ""))
    run = p.add_run(f"Kundennummer: {kd}  ·  Abrechnungszeitraum: {von} – {bis}")
    run.font.size = Pt(9)

    spacer()

    # ── Seite 3: Wohnungsnutzer ───────────────────────────────────────────────
    add_section_heading("Seite 3  –  Wohnungsnutzer")
    nutzer = data.get("wohnungsnutzer", [])
    add_table(
        ["Whg.", "Lage", "Eigentümer", "Nutzer", "Fläche m²", "Beh. m²", "Pers.", "Vorausz. €"],
        [
            [
                _s(r.get("whg_nr")), _s(r.get("lage")),
                _s(r.get("eigentuemer")), _s(r.get("nutzer")),
                _n(r.get("wohnflaeche")), _n(r.get("beheizte_wohnflaeche")),
                _s(r.get("personen") if r.get("personen") else ""),
                _n(r.get("vorauszahlungen")),
            ]
            for r in nutzer
        ],
        [Cm(1.2), Cm(1.8), Cm(3.0), Cm(3.0), Cm(1.6), Cm(1.6), Cm(0.9), Cm(2.8)],
    )

    # ── Seite 1: Nebenkosten ──────────────────────────────────────────────────
    spacer()
    add_section_heading("Seite 1  –  Nebenkosten")
    _nk_h = ["Kostenart", "Lieferant", "Nr.", "Rechnungsdatum", "Betrag Brutto €",
             "Gesamtbetrag €", "Anteil AK €", "Whg. Nr."]
    _nk_w = [Cm(3.2), Cm(2.8), Cm(1.0), Cm(1.8), Cm(1.8), Cm(1.8), Cm(1.5), Cm(1.0)]

    def nk_row(r):
        return [
            _s(r.get("kategorie")), _s(r.get("lieferant")), _s(r.get("nr")),
            _d(r.get("rechnungsdatum")), _n(r.get("betrag_brutto")),
            _n(r.get("gesamtbetrag")), _n(r.get("anteil_arbeitskosten")),
            _s(r.get("whg_nr")),
        ]

    add_table(_nk_h, [nk_row(r) for r in data.get("nebenkosten", [])], _nk_w)

    # ── Seite 1: Wasserkosten ─────────────────────────────────────────────────
    spacer()
    add_section_heading("Seite 1  –  Wasserkosten")
    add_table(_nk_h, [nk_row(r) for r in data.get("wasserkosten", [])], _nk_w)

    doc.add_page_break()

    # ── Seite 2: Stammdaten Heizanlage ────────────────────────────────────────
    add_section_heading("Seite 2  –  Stammdaten Heizanlage")
    sh = data.get("stammdaten_heizung", {})
    add_table(
        ["Parameter", "Wert"],
        [
            ["Brennstoffart",                     _s(sh.get("brennstoffart"))],
            ["Mengeneinheit",                     _s(sh.get("mengeneinheit"))],
            ["Heizwert (kWh/ME)",                 _n(sh.get("heizwert"))],
            ["Emissionsfaktor (kg CO₂/kWh)",      _n(sh.get("emissionsfaktor")) or "–"],
            ["Warmwasser Grundkostenanteil %",     _n(sh.get("warmwasser_grundkostenanteil"))],
            ["Heizung Grundkostenanteil %",        _n(sh.get("heizung_grundkostenanteil"))],
            ["Warmwassertemperatur °C",            _n(sh.get("warmwasser_temperatur"), 0)],
        ],
        [Cm(8), Cm(4)],
    )

    # ── Seite 2: Brennstoffkosten ─────────────────────────────────────────────
    spacer()
    add_section_heading("Seite 2  –  Brennstoffkosten")
    bk  = data.get("brennstoffkosten", {})
    me  = sh.get("mengeneinheit", "ME")
    bk_rows = [
        [
            f"Vorjahresbestand  ({_d(bk.get('vorjahresbestand_datum'))})",
            _n(bk.get("vorjahresbestand_menge"), 0),
            _n(bk.get("vorjahresbestand_betrag")),
        ],
    ]
    total_preis = float(bk.get("vorjahresbestand_betrag") or 0)
    for i, e in enumerate(bk.get("einkaufe", []), 1):
        preis = float(e.get("betrag_brutto") or 0)
        total_preis += preis
        bk_rows.append([
            f"Einkauf {i}  ({_d(e.get('datum'))})",
            _n(e.get("menge"), 0),
            _n(preis),
        ])
    vj       = float(bk.get("vorjahresbestand_menge") or 0)
    je       = float(bk.get("jahresendbestand_menge") or 0)
    zugaenge = sum(float(e.get("menge") or 0) for e in bk.get("einkaufe", []))
    verbrauch = vj + zugaenge - je
    bk_rows.append(["Jahresendbestand", _n(je, 0), ""])
    bk_rows.append(["Gesamtverbrauch", _n(verbrauch, 0), _n(total_preis)])
    add_table(
        ["Position", f"Menge ({me})", "Betrag €"],
        bk_rows,
        [Cm(7), Cm(3), Cm(3)],
    )

    # ── Seite 2: Heiznebenkosten ──────────────────────────────────────────────
    spacer()
    add_section_heading("Seite 2  –  Heiznebenkosten")
    add_table(
        ["Kostenart", "Datum", "Nr.", "Betrag Brutto €", "Gesamtbetrag €", "Anteil AK €"],
        [
            [
                _s(r.get("kategorie")), _d(r.get("datum")), _s(r.get("nr")),
                _n(r.get("betrag_brutto")), _n(r.get("gesamtbetrag")),
                _n(r.get("anteil_arbeitskosten")),
            ]
            for r in data.get("heiznebenkosten", [])
        ],
        [Cm(4.5), Cm(1.8), Cm(1.2), Cm(2.3), Cm(2.3), Cm(2.2)],
    )

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ─── PDF Export ───────────────────────────────────────────────────────────────

def generate_pdf(data: dict, year: int) -> BytesIO:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import (
        SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak,
    )
    from reportlab.lib.enums import TA_CENTER

    PRIMARY = colors.Color(26 / 255, 122 / 255, 94 / 255)   # #1A7A5E
    LIGHT   = colors.Color(230 / 255, 244 / 255, 239 / 255) # #E6F4EF
    GRID    = colors.Color(0.75, 0.75, 0.75)

    buf = BytesIO()
    doc_obj = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=2 * cm,  rightMargin=2 * cm,
        topMargin=2.5 * cm, bottomMargin=2 * cm,
        title=f"Witter Heizkostenabrechnung {year}",
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "wt_title", parent=styles["Normal"],
        fontSize=16, textColor=PRIMARY, fontName="Helvetica-Bold",
        spaceBefore=0, spaceAfter=4, alignment=TA_CENTER,
    )
    sub_style = ParagraphStyle(
        "wt_sub", parent=styles["Normal"],
        fontSize=9, spaceAfter=10, alignment=TA_CENTER,
    )
    h1_style = ParagraphStyle(
        "wt_h1", parent=styles["Normal"],
        fontSize=10, textColor=PRIMARY, fontName="Helvetica-Bold",
        spaceBefore=12, spaceAfter=3,
        borderPadding=(0, 0, 2, 0),
    )

    def pdf_table(headers: list, rows: list, widths: list) -> Table:
        tdata = [headers] + rows
        t = Table(tdata, colWidths=widths, repeatRows=1)
        ts = [
            ("BACKGROUND",   (0, 0), (-1, 0), PRIMARY),
            ("TEXTCOLOR",    (0, 0), (-1, 0), colors.white),
            ("FONTNAME",     (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE",     (0, 0), (-1, -1), 8),
            ("TOPPADDING",   (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING",(0, 0), (-1, -1), 3),
            ("LEFTPADDING",  (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("GRID",         (0, 0), (-1, -1), 0.3, GRID),
            ("VALIGN",       (0, 0), (-1, -1), "MIDDLE"),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, LIGHT]),
        ]
        t.setStyle(TableStyle(ts))
        return t

    story: list = []

    # ── Title ─────────────────────────────────────────────────────────────────
    story.append(Paragraph(f"Witter Heizkostenabrechnung {year}", title_style))
    kd  = data.get("kundennummer", "")
    von = _d(data.get("von", ""))
    bis = _d(data.get("bis", ""))
    story.append(Paragraph(f"Kundennummer: {kd}  ·  Abrechnungszeitraum: {von} – {bis}", sub_style))

    # ── Wohnungsnutzer ────────────────────────────────────────────────────────
    story.append(Paragraph("Wohnungsnutzer", h1_style))
    nutzer = data.get("wohnungsnutzer", [])
    story.append(pdf_table(
        ["Whg.", "Lage", "Eigentümer", "Nutzer", "Fläche\nm²", "Beh.\nm²", "Pers.", "Vorausz.\n€"],
        [
            [
                _s(r.get("whg_nr")), _s(r.get("lage")),
                _s(r.get("eigentuemer")), _s(r.get("nutzer")),
                _n(r.get("wohnflaeche")), _n(r.get("beheizte_wohnflaeche")),
                _s(r.get("personen") if r.get("personen") else ""),
                _n(r.get("vorauszahlungen")),
            ]
            for r in nutzer
        ],
        [1.2*cm, 1.8*cm, 3.0*cm, 3.0*cm, 1.6*cm, 1.6*cm, 0.9*cm, 2.8*cm],
    ))

    # ── Nebenkosten ───────────────────────────────────────────────────────────
    story.append(Spacer(1, 0.4 * cm))
    story.append(Paragraph("Nebenkosten", h1_style))
    _nk_h  = ["Kostenart", "Lieferant", "Nr.", "Datum", "Betrag\nBrutto €",
              "Gesamt-\nbetrag €", "Anteil\nAK €", "Whg.\nNr."]
    _nk_cw = [3.2*cm, 2.8*cm, 1.0*cm, 1.8*cm, 1.8*cm, 1.8*cm, 1.5*cm, 1.0*cm]

    def nk_row(r):
        return [
            _s(r.get("kategorie")), _s(r.get("lieferant")), _s(r.get("nr")),
            _d(r.get("rechnungsdatum")), _n(r.get("betrag_brutto")),
            _n(r.get("gesamtbetrag")), _n(r.get("anteil_arbeitskosten")),
            _s(r.get("whg_nr")),
        ]

    story.append(pdf_table(_nk_h, [nk_row(r) for r in data.get("nebenkosten", [])], _nk_cw))

    # ── Wasserkosten ──────────────────────────────────────────────────────────
    story.append(Spacer(1, 0.4 * cm))
    story.append(Paragraph("Wasserkosten", h1_style))
    story.append(pdf_table(_nk_h, [nk_row(r) for r in data.get("wasserkosten", [])], _nk_cw))

    story.append(PageBreak())

    # ── Stammdaten Heizanlage ─────────────────────────────────────────────────
    story.append(Paragraph("Stammdaten Heizanlage", h1_style))
    sh = data.get("stammdaten_heizung", {})
    story.append(pdf_table(
        ["Parameter", "Wert"],
        [
            ["Brennstoffart",                    _s(sh.get("brennstoffart"))],
            ["Mengeneinheit",                    _s(sh.get("mengeneinheit"))],
            ["Heizwert (kWh/ME)",                _n(sh.get("heizwert"))],
            ["Emissionsfaktor (kg CO\u2082/kWh)",_n(sh.get("emissionsfaktor")) or "\u2013"],
            ["Warmwasser Grundkostenanteil %",   _n(sh.get("warmwasser_grundkostenanteil"))],
            ["Heizung Grundkostenanteil %",      _n(sh.get("heizung_grundkostenanteil"))],
            ["Warmwassertemperatur \u00b0C",     _n(sh.get("warmwasser_temperatur"), 0)],
        ],
        [8 * cm, 4 * cm],
    ))

    # ── Brennstoffkosten ──────────────────────────────────────────────────────
    story.append(Spacer(1, 0.4 * cm))
    story.append(Paragraph("Brennstoffkosten", h1_style))
    bk    = data.get("brennstoffkosten", {})
    me    = sh.get("mengeneinheit", "ME")
    bk_rows: list = [
        [
            f"Vorjahresbestand  ({_d(bk.get('vorjahresbestand_datum'))})",
            _n(bk.get("vorjahresbestand_menge"), 0),
            _n(bk.get("vorjahresbestand_betrag")),
        ],
    ]
    total_preis = float(bk.get("vorjahresbestand_betrag") or 0)
    for i, e in enumerate(bk.get("einkaufe", []), 1):
        preis = float(e.get("betrag_brutto") or 0)
        total_preis += preis
        bk_rows.append([f"Einkauf {i}  ({_d(e.get('datum'))})", _n(e.get("menge"), 0), _n(preis)])
    vj        = float(bk.get("vorjahresbestand_menge") or 0)
    je        = float(bk.get("jahresendbestand_menge") or 0)
    zugaenge  = sum(float(e.get("menge") or 0) for e in bk.get("einkaufe", []))
    verbrauch = vj + zugaenge - je
    bk_rows.append(["Jahresendbestand",  _n(je, 0),        ""])
    bk_rows.append(["Gesamtverbrauch",   _n(verbrauch, 0), _n(total_preis)])
    story.append(pdf_table(
        ["Position", f"Menge ({me})", "Betrag \u20ac"],
        bk_rows,
        [8 * cm, 3 * cm, 3 * cm],
    ))

    # ── Heiznebenkosten ───────────────────────────────────────────────────────
    story.append(Spacer(1, 0.4 * cm))
    story.append(Paragraph("Heiznebenkosten", h1_style))
    story.append(pdf_table(
        ["Kostenart", "Datum", "Nr.", "Betrag\nBrutto \u20ac", "Gesamt-\nbetrag \u20ac", "Anteil\nAK \u20ac"],
        [
            [
                _s(r.get("kategorie")), _d(r.get("datum")), _s(r.get("nr")),
                _n(r.get("betrag_brutto")), _n(r.get("gesamtbetrag")),
                _n(r.get("anteil_arbeitskosten")),
            ]
            for r in data.get("heiznebenkosten", [])
        ],
        [4.5*cm, 1.8*cm, 1.2*cm, 2.3*cm, 2.3*cm, 2.3*cm],
    ))

    doc_obj.build(story)
    buf.seek(0)
    return buf
