"""
Eigentümerabrechnung – DOCX Export Service

Generates an owner billing letter (WEG annual settlement) for a single
Wohnungseinheit as a Word (.docx) document.
"""

from io import BytesIO
from datetime import date as _date

# ─── Static bank balance data per year ────────────────────────────────────────

KONTENENTWICKLUNG = {
    2024: {
        "giro":      [("29.12.2023", 3_527.48), ("30.12.2024", 3_333.80)],
        "geldmarkt": [("29.12.2023", 8_067.85), ("30.12.2024", 10_274.95)],
    },
    2025: {
        "giro":      [("30.12.2024", 3_333.80), ("30.12.2025", 3_970.48)],
        "geldmarkt": [("30.12.2024", 10_281.93), ("30.12.2025", 12_677.11)],
    },
}

# ─── Colour palette ───────────────────────────────────────────────────────────
# Deep slate-teal header, very light tint for alternating rows, a warm amber
# accent for the result row.

_PRIMARY_HEX  = "1F4E79"   # deep navy-blue  (header bg)
_ACCENT_HEX   = "2E75B6"   # mid-blue        (subject underline, section labels)
_ALT_HEX      = "EBF3FB"   # pale blue tint  (odd data rows)
_RESULT_HEX   = "FFF2CC"   # pale amber      (Nachzahlung / Rückzahlung highlight)
_BORDER_HEX   = "BDD7EE"   # light-blue rule (table inner lines)
_GRAY_HEX     = "666666"   # mid-gray        (secondary text)


# ─── Formatting helpers ────────────────────────────────────────────────────────

def _n(val, dec: int = 2) -> str:
    """Format number German-style: 1.234,56"""
    if val is None:
        return "ausstehend"
    try:
        s = f"{float(val):,.{dec}f}"
        return s.replace(",", "\x00").replace(".", ",").replace("\x00", ".")
    except (TypeError, ValueError):
        return "ausstehend"


def _eur(val) -> str:
    if val is None:
        return "ausstehend"
    return _n(val) + " €"


# ─── DOCX generator ───────────────────────────────────────────────────────────

def generate_owner_docx(owner: dict, year: int) -> BytesIO:
    """Generate a single Eigentümerabrechnung DOCX for `owner` and `year`."""
    from docx import Document
    from docx.shared import Pt, Cm, Emu, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from lxml import etree

    PRIMARY_RGB = RGBColor(0x1F, 0x4E, 0x79)
    ACCENT_RGB  = RGBColor(0x2E, 0x75, 0xB6)
    WHITE_RGB   = RGBColor(0xFF, 0xFF, 0xFF)
    GRAY_RGB    = RGBColor(0x66, 0x66, 0x66)
    DARK_RGB    = RGBColor(0x1A, 0x1A, 0x2E)

    FONT_BODY   = "Calibri"
    FONT_HEAD   = "Calibri"

    doc = Document()

    # ── Page setup: A4, narrow margins (DIN 5008 style) ───────────────────────
    sec = doc.sections[0]
    sec.page_width    = Cm(21)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.0)
    sec.top_margin    = Cm(2.0)
    sec.bottom_margin = Cm(2.0)

    # ── Low-level XML helpers ──────────────────────────────────────────────────

    def set_cell_bg(cell, hex_color: str):
        tcPr = cell._tc.get_or_add_tcPr()
        for e in tcPr.findall(qn("w:shd")):
            tcPr.remove(e)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color)
        tcPr.append(shd)

    def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
        """Set individual cell borders. Pass hex colour string or None to skip."""
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = tcPr.find(qn("w:tcBorders"))
        if tcBorders is None:
            tcBorders = OxmlElement("w:tcBorders")
            tcPr.append(tcBorders)
        sides = {"top": top, "bottom": bottom, "left": left, "right": right}
        for side, color in sides.items():
            existing = tcBorders.find(qn(f"w:{side}"))
            if existing is not None:
                tcBorders.remove(existing)
            el = OxmlElement(f"w:{side}")
            if color:
                el.set(qn("w:val"),   "single")
                el.set(qn("w:sz"),    "4")
                el.set(qn("w:space"), "0")
                el.set(qn("w:color"), color)
            else:
                el.set(qn("w:val"), "nil")
            tcBorders.append(el)

    def remove_all_borders(cell):
        set_cell_border(cell, top=None, bottom=None, left=None, right=None)

    def set_para_bottom_border(p, color_hex: str, sz: str = "12"):
        pPr  = p._p.get_or_add_pPr()
        pBdr = pPr.find(qn("w:pBdr"))
        if pBdr is None:
            pBdr = OxmlElement("w:pBdr")
            pPr.append(pBdr)
        for e in pBdr.findall(qn("w:bottom")):
            pBdr.remove(e)
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    sz)
        bot.set(qn("w:space"), "4")
        bot.set(qn("w:color"), color_hex)
        pBdr.append(bot)

    # ── Paragraph / run factory ────────────────────────────────────────────────

    def add_para(text="", bold=False, italic=False, size_pt=10.5,
                 align=None, color_rgb=None,
                 space_before_pt=0, space_after_pt=0, font=FONT_BODY):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before_pt)
        p.paragraph_format.space_after  = Pt(space_after_pt)
        if align:
            p.alignment = align
        if text:
            r = p.add_run(text)
            r.font.name  = font
            r.bold       = bold
            r.italic     = italic
            r.font.size  = Pt(size_pt)
            if color_rgb:
                r.font.color.rgb = color_rgb
        return p

    def set_cell_text(cell, text="", bold=False, italic=False, size_pt=9.5,
                      align=WD_ALIGN_PARAGRAPH.LEFT, color_rgb=None, font=FONT_BODY):
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = align
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        if text:
            r = p.add_run(text)
            r.font.name  = font
            r.bold       = bold
            r.italic     = italic
            r.font.size  = Pt(size_pt)
            if color_rgb:
                r.font.color.rgb = color_rgb

    # ══════════════════════════════════════════════════════════════════════════
    # LETTERHEAD
    # ══════════════════════════════════════════════════════════════════════════

    # ── Organisation name (large, primary colour) ──────────────────────────────
    p_org = add_para(
        "WEG Tulpenstraße 31",
        bold=True, size_pt=14, color_rgb=PRIMARY_RGB,
        space_after_pt=0, font=FONT_HEAD,
    )
    p_org.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # ── Sub-line: address + contact ────────────────────────────────────────────
    p_sub_addr = add_para(
        "86567 Hilgertshausen  ·  Verwaltung: WEG-Gemeinschaft",
        size_pt=8.5, color_rgb=GRAY_RGB, space_after_pt=2,
    )
    set_para_bottom_border(p_sub_addr, _PRIMARY_HEX, sz="16")

    add_para(space_after_pt=4)  # visual gap below header rule

    # ══════════════════════════════════════════════════════════════════════════
    # ADDRESS BLOCK  +  DATE (two-column via table)
    # ══════════════════════════════════════════════════════════════════════════

    addr_tbl = doc.add_table(rows=1, cols=2)
    left_cell  = addr_tbl.rows[0].cells[0]
    right_cell = addr_tbl.rows[0].cells[1]
    left_cell.width  = Cm(10)
    right_cell.width = Cm(5.5)
    remove_all_borders(left_cell)
    remove_all_borders(right_cell)

    # Sender hint (small, above address block)
    p_hint = left_cell.add_paragraph()
    p_hint.paragraph_format.space_before = Pt(0)
    p_hint.paragraph_format.space_after  = Pt(2)
    rh = p_hint.add_run("WEG Tulpenstr. 31 · 86567 Hilgertshausen")
    rh.font.name  = FONT_BODY
    rh.font.size  = Pt(7.5)
    rh.font.color.rgb = GRAY_RGB
    rh.underline  = True

    for line in [owner.get("name", ""), owner.get("address1", ""), owner.get("city", "")]:
        p_a = left_cell.add_paragraph()
        p_a.paragraph_format.space_before = Pt(0)
        p_a.paragraph_format.space_after  = Pt(0)
        ra = p_a.add_run(line)
        ra.font.name = FONT_BODY
        ra.font.size = Pt(10.5)
        ra.bold = (line == owner.get("name", ""))

    # Date – right-aligned in right cell
    today = _date.today().strftime("%d.%m.%Y")
    p_dt = right_cell.add_paragraph()
    p_dt.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_dt.paragraph_format.space_before = Pt(0)
    p_dt.paragraph_format.space_after  = Pt(0)
    rd = p_dt.add_run(today)
    rd.font.name  = FONT_BODY
    rd.font.size  = Pt(10)
    rd.font.color.rgb = GRAY_RGB

    add_para(space_after_pt=6)

    # ══════════════════════════════════════════════════════════════════════════
    # SUBJECT LINE
    # ══════════════════════════════════════════════════════════════════════════

    unit_id = owner.get("unit_id", "")
    p_subj = add_para(
        f"Jahresabrechnung {year}  ·  {unit_id}",
        bold=True, size_pt=13, color_rgb=PRIMARY_RGB,
        space_before_pt=2, space_after_pt=2, font=FONT_HEAD,
    )
    set_para_bottom_border(p_subj, _ACCENT_HEX, sz="8")

    add_para(space_after_pt=4)

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION LABEL: Ausgaben
    # ══════════════════════════════════════════════════════════════════════════

    p_lbl = add_para(
        "Ausgaben und Kostenverteilung",
        bold=True, size_pt=10, color_rgb=ACCENT_RGB,
        space_after_pt=3,
    )

    # ══════════════════════════════════════════════════════════════════════════
    # COST TABLE
    # ══════════════════════════════════════════════════════════════════════════

    col_widths = [Cm(5.8), Cm(3.8), Cm(2.1), Cm(2.6), Cm(2.8)]
    tbl = doc.add_table(rows=0, cols=5)

    def add_tbl_row(cells_data, bg=None, bold=False, italic=False,
                    top_border=None, bottom_border=None,
                    text_color=None, size_pt=9.5):
        row = tbl.add_row()
        is_header = (bg == _PRIMARY_HEX)
        for ci, (width, text) in enumerate(zip(col_widths, cells_data)):
            row.cells[ci].width = width
            align = WD_ALIGN_PARAGRAPH.RIGHT if ci >= 2 else WD_ALIGN_PARAGRAPH.LEFT
            fg = text_color or (WHITE_RGB if is_header else DARK_RGB)
            set_cell_text(row.cells[ci], text,
                          bold=bold, italic=italic, size_pt=size_pt,
                          align=align, color_rgb=fg)
            if bg:
                set_cell_bg(row.cells[ci], bg)
            # Remove outer vertical borders; keep subtle horizontal lines
            set_cell_border(
                row.cells[ci],
                top=top_border,
                bottom=bottom_border,
                left=None,
                right=None,
            )

    # ── Header row ────────────────────────────────────────────────────────────
    add_tbl_row(
        ["Kostenkonto", "Umlageart", "MEA", "Ges. Betrag", "Ant. Betrag"],
        bg=_PRIMARY_HEX, bold=True, size_pt=9,
    )

    # ── Data rows (alternating) ────────────────────────────────────────────────
    positionen = owner.get("positionen", [])
    for idx, pos in enumerate(positionen):
        mea = pos.get("umlage_anteilig")
        mea_str = _n(mea) if mea is not None else "gem. BetrKV"
        bg = _ALT_HEX if idx % 2 == 1 else None
        add_tbl_row(
            [
                pos.get("kostenkonto", ""),
                pos.get("umlageart", ""),
                mea_str,
                _n(pos.get("ges_betrag")),
                _eur(pos.get("betrag_ant")),
            ],
            bg=bg,
            bottom_border=_BORDER_HEX,
        )

    # ── Thin separator ────────────────────────────────────────────────────────
    add_tbl_row(["", "", "", "", ""],
                top_border=_PRIMARY_HEX, bottom_border=_PRIMARY_HEX)

    # ── Gesamt ────────────────────────────────────────────────────────────────
    gesamt_ges = None
    try:
        gesamt_ges = sum(
            p["ges_betrag"] for p in positionen if p.get("ges_betrag") is not None
        )
        if any(p.get("ges_betrag") is None for p in positionen):
            gesamt_ges = None
    except Exception:
        gesamt_ges = None

    add_tbl_row(
        ["Gesamtkosten", "", "", _n(gesamt_ges), _eur(owner.get("gesamt"))],
        bold=True, bottom_border=_BORDER_HEX,
    )

    # ── Vorauszahlung ─────────────────────────────────────────────────────────
    add_tbl_row(
        ["Abzüglich Vorauszahlung", "", "", "", _eur(owner.get("vorauszahlung"))],
        italic=True, bottom_border=_BORDER_HEX,
        text_color=GRAY_RGB,
    )

    # ── Nachzahlung / Rückzahlung (amber highlight) ───────────────────────────
    nachzahlung = owner.get("nachzahlung")
    nz_label = "Nachzahlung" if nachzahlung is None or nachzahlung >= 0 else "Rückzahlung"
    add_tbl_row(
        [nz_label, "", "", "", _eur(nachzahlung)],
        bg=_RESULT_HEX, bold=True, size_pt=10,
        top_border=_PRIMARY_HEX, bottom_border=_PRIMARY_HEX,
    )

    add_para(space_after_pt=4)

    # ── Footnote ───────────────────────────────────────────────────────────────
    p_fn = add_para(
        "¹ Gesamtkosten Betriebskostenabrechnung gem. BetrKV – Details aus der "
        "Heiz- und Nebenkostenabrechnung der Fa. Witter.",
        size_pt=7.5, color_rgb=GRAY_RGB, space_after_pt=6,
    )

    # ══════════════════════════════════════════════════════════════════════════
    # KONTENENTWICKLUNG
    # ══════════════════════════════════════════════════════════════════════════

    ke = KONTENENTWICKLUNG.get(year)
    if ke:
        p_ke_lbl = add_para(
            "Kontenentwicklung",
            bold=True, size_pt=10, color_rgb=ACCENT_RGB,
            space_before_pt=2, space_after_pt=3,
        )

        # Small 3-column table: Konto | Datum | Betrag
        ktbl = doc.add_table(rows=0, cols=3)
        k_widths = [Cm(4.5), Cm(3.5), Cm(3)]

        def add_ktbl_row(cols_data, bold=False, bg=None):
            r = ktbl.add_row()
            for ci, (w, txt) in enumerate(zip(k_widths, cols_data)):
                r.cells[ci].width = w
                set_cell_text(
                    r.cells[ci], txt,
                    bold=bold, size_pt=9.5,
                    align=WD_ALIGN_PARAGRAPH.RIGHT if ci == 2 else WD_ALIGN_PARAGRAPH.LEFT,
                    color_rgb=WHITE_RGB if bg == _PRIMARY_HEX else DARK_RGB,
                )
                if bg:
                    set_cell_bg(r.cells[ci], bg)
                else:
                    set_cell_border(r.cells[ci], bottom=_BORDER_HEX)

        add_ktbl_row(["Konto", "Stand", "Betrag"], bold=True, bg=_PRIMARY_HEX)
        prev_year = year - 1
        add_ktbl_row(["Girokonto", ke["giro"][0][0], _eur(ke["giro"][0][1])], bg=_ALT_HEX)
        add_ktbl_row(["Girokonto", ke["giro"][1][0], _eur(ke["giro"][1][1])])
        add_ktbl_row(["Geldmarktkonto", ke["geldmarkt"][0][0], _eur(ke["geldmarkt"][0][1])], bg=_ALT_HEX)
        add_ktbl_row(["Geldmarktkonto", ke["geldmarkt"][1][0], _eur(ke["geldmarkt"][1][1])])

        add_para(space_after_pt=6)

    # ══════════════════════════════════════════════════════════════════════════
    # CLOSING / ANLAGEN
    # ══════════════════════════════════════════════════════════════════════════

    add_para(
        "Mit freundlichen Grüßen",
        size_pt=10, space_before_pt=2, space_after_pt=12,
    )
    add_para(
        "WEG-Gemeinschaft Tulpenstraße 31",
        bold=True, size_pt=10, color_rgb=PRIMARY_RGB,
    )

    add_para(space_after_pt=8)

    p_anl = add_para("Anlage:", bold=True, size_pt=9, color_rgb=GRAY_RGB)
    add_para(
        "Witter Heiz- und Nebenkostenabrechnung",
        size_pt=9, color_rgb=GRAY_RGB,
    )

    # ── Write to buffer ────────────────────────────────────────────────────────
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
